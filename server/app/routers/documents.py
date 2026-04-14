import io

import database
import docx
import models
from fastapi import APIRouter, Depends, File, HTTPException, Response, UploadFile
from pydantic import BaseModel
from services import ai_engine, parser
from sqlalchemy.orm import Session

router = APIRouter(prefix="/api/documents", tags=["Documents"])


# Модель для получения сообщения из Frontend
class ChatMessage(BaseModel):
    message: str


class GrantCheckRequest(BaseModel):
    grant_name: str = "Грант Комитета Науки (стандартный)"


@router.post("/upload-and-analyze/")
async def upload_and_analyze(
    file: UploadFile = File(...), db: Session = Depends(database.get_db)
):
    # 1. Чтение файла
    content = await file.read()

    # 2. Извлечение текста
    try:
        text = parser.parse_document(content, file.filename)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))

    # 3. Сохранение информации о документе в БД
    file_ext = file.filename.split(".")[-1]
    new_doc = models.Document(
        filename=file.filename, file_format=file_ext, full_text=text
    )
    # Примечание: пока без привязки к owner_id, добавим при реализации авторизации
    db.add(new_doc)
    db.commit()
    db.refresh(new_doc)

    # 4. Запуск AI-анализа (параллельно для скорости)
    # В идеале использовать asyncio.gather, но для простоты прототипа вызовем последовательно
    analysis_data = await ai_engine.analyze_tz_content(text)
    score = await ai_engine.score_tz(text)
    improvements_data = await ai_engine.generate_improvements(text)

    # 5. Сохранение результатов анализа в БД
    result = models.AnalysisResult(
        document_id=new_doc.id,
        score=score,
        errors_found=analysis_data,
        recommendations=improvements_data.get("recommendations", []),
        structure_data=improvements_data.get("recommended_structure", []),
    )
    db.add(result)
    db.commit()

    # 6. Возврат итогового отчета на Frontend
    return {
        "message": "Анализ успешно завершен",
        "document_id": new_doc.id,
        "score": score,
        "errors": analysis_data,
        "improvements": improvements_data,
    }


# 1. ЭНДПОИНТ: Получение истории всех документов для Дашборда
@router.get("/")
def get_all_documents(db: Session = Depends(database.get_db)):
    docs = db.query(models.Document).all()
    result = []
    for doc in docs:
        result.append(
            {
                "id": doc.id,
                "filename": doc.filename,
                "upload_time": doc.upload_time,
                "score": doc.analysis.score if doc.analysis else None,
            }
        )
    return result


# 2. ЭНДПОИНТ: Получение результатов анализа конкретного документа
@router.get("/{document_id}/analysis")
def get_document_analysis(document_id: int, db: Session = Depends(database.get_db)):
    analysis = (
        db.query(models.AnalysisResult)
        .filter(models.AnalysisResult.document_id == document_id)
        .first()
    )
    if not analysis:
        raise HTTPException(
            status_code=404, detail="Анализ для данного документа не найден"
        )

    return {
        "score": analysis.score,
        "errors": analysis.errors_found,
        "improvements": analysis.recommendations,
        "structure": analysis.structure_data,
    }


# 3. ЭНДПОИНТ: Чат с AI-ассистентом по документу
@router.post("/{document_id}/chat")
async def chat_about_document(
    document_id: int, chat_data: ChatMessage, db: Session = Depends(database.get_db)
):
    analysis = (
        db.query(models.AnalysisResult)
        .filter(models.AnalysisResult.document_id == document_id)
        .first()
    )
    if not analysis:
        raise HTTPException(status_code=404, detail="Документ не найден")

    # Превращаем JSON-данные из БД в строку для контекста ИИ
    recs_str = str(analysis.recommendations)
    struct_str = str(analysis.structure_data)

    # Запрашиваем ответ у LLM
    bot_reply = await ai_engine.chat_with_assistant(
        chat_data.message, struct_str, recs_str
    )

    # Сохраняем историю в БД
    new_msg_user = models.ChatHistory(
        document_id=document_id, message=chat_data.message, is_bot=False
    )
    new_msg_bot = models.ChatHistory(
        document_id=document_id, message=bot_reply, is_bot=True
    )
    db.add(new_msg_user)
    db.add(new_msg_bot)
    db.commit()

    return {"reply": bot_reply}


# 4. ЭНДПОИНТ: Скачивание улучшенной структуры ТЗ в виде текстового файла
# @router.get("/{document_id}/download")
# def download_improved_tz(document_id: int, db: Session = Depends(database.get_db)):
#     doc = db.query(models.Document).filter(models.Document.id == document_id).first()
#     analysis = (
#         db.query(models.AnalysisResult)
#         .filter(models.AnalysisResult.document_id == document_id)
#         .first()
#     )

#     if not analysis or not doc:
#         raise HTTPException(status_code=404, detail="Данные не найдены")

#     # Формируем красивый текст для скачивания
#     content = f"=== УЛУЧШЕННАЯ СТРУКТУРА ТЗ: {doc.filename} ===\n\n"

#     if (
#         isinstance(analysis.structure_data, dict)
#         and "recommended_structure" in analysis.structure_data
#     ):
#         for item in analysis.structure_data["recommended_structure"]:
#             content += f"- {item}\n"
#     elif isinstance(analysis.structure_data, list):
#         for item in analysis.structure_data:
#             content += f"- {item}\n"

#     content += "\n=== КЛЮЧЕВЫЕ РЕКОМЕНДАЦИИ ===\n\n"
#     if isinstance(analysis.recommendations, list):
#         for rec in analysis.recommendations:
#             content += f"* {rec}\n"

#     safe_filename = f"improved_tz_{doc.id}.txt"  # Только латиница + ID
#     # Возвращаем файл напрямую в браузер пользователя
#     return Response(
#         content=content,
#         media_type="text/plain",
#         headers={
#             "Content-Disposition": f"attachment; filename=\"{safe_filename}\"; filename*=UTF-8''{safe_filename}"
#         },
#     )


@router.get("/{document_id}/export-docx")
def export_improved_tz_docx(document_id: int, db: Session = Depends(database.get_db)):
    doc_record = (
        db.query(models.Document).filter(models.Document.id == document_id).first()
    )
    analysis = (
        db.query(models.AnalysisResult)
        .filter(models.AnalysisResult.document_id == document_id)
        .first()
    )

    if not analysis or not doc_record:
        raise HTTPException(status_code=404, detail="Данные не найдены")

    # 1. Создаем новый Word-документ
    document = docx.Document()

    # 2. Добавляем красивый заголовок
    document.add_heading("Улучшенное Техническое Задание", 0)
    document.add_paragraph(
        f"Сгенерировано AI-ассистентом SHAM на основе файла: {doc_record.filename}\n"
    )

    # 3. Секция: Рекомендуемая структура
    document.add_heading("1. Рекомендуемая структура ТЗ", level=1)

    struct_data = analysis.structure_data
    if isinstance(struct_data, dict) and "recommended_structure" in struct_data:
        for item in struct_data["recommended_structure"]:
            document.add_paragraph(item, style="List Bullet")
    elif isinstance(struct_data, list):
        for item in struct_data:
            document.add_paragraph(item, style="List Bullet")
    else:
        document.add_paragraph("Структура не сгенерирована.")

    # 4. Секция: Рекомендации по доработке
    document.add_heading("2. Рекомендации по доработке", level=1)

    if isinstance(analysis.recommendations, list):
        for rec in analysis.recommendations:
            document.add_paragraph(rec, style="List Bullet")
    else:
        document.add_paragraph("Рекомендации отсутствуют.")

    # 5. Сохраняем документ в оперативную память (BytesIO)
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)  # Возвращаем курсор в начало файла

    # 6. Отправляем файл пользователю
    # Указываем правильный MIME-тип для .docx
    headers = {
        "Content-Disposition": f"attachment; filename=improved_tz_{document_id}.docx"
    }

    return Response(
        content=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )


# Новый эндпоинт
@router.post("/{document_id}/check-grant")
async def check_grant_compliance_endpoint(
    document_id: int, request: GrantCheckRequest, db: Session = Depends(database.get_db)
):
    # 1. Ищем документ
    doc = db.query(models.Document).filter(models.Document.id == document_id).first()
    if not doc or not doc.full_text:
        raise HTTPException(status_code=404, detail="Документ или его текст не найден")

    # 2. Отправляем текст в Groq
    grant_analysis = await ai_engine.check_grant_compliance(
        doc.full_text, request.grant_name
    )

    # В рамках хакатона можно не сохранять это в БД, а отдавать прямо на фронт (on-demand генерация)
    return grant_analysis


@router.post("/{document_id}/compare-template")
async def compare_template_endpoint(
    document_id: int, db: Session = Depends(database.get_db)
):
    # Находим документ и извлекаем его полный текст (который мы добавили на предыдущем шаге)
    doc = db.query(models.Document).filter(models.Document.id == document_id).first()
    if not doc or not doc.full_text:
        raise HTTPException(status_code=404, detail="Документ или его текст не найден")

    # Запускаем сравнение
    comparison_result = await ai_engine.compare_with_template(doc.full_text)

    return comparison_result


@router.get("/{document_id}/metrics")
async def get_document_metrics(
    document_id: int, db: Session = Depends(database.get_db)
):
    # Ищем документ и его текст
    doc = db.query(models.Document).filter(models.Document.id == document_id).first()
    if not doc or not doc.full_text:
        raise HTTPException(status_code=404, detail="Документ или его текст не найден")

    # Запускаем извлечение метрик
    metrics_data = await ai_engine.extract_metrics(doc.full_text)

    return metrics_data
