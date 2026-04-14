import io

import docx
from PyPDF2 import PdfReader

# def extract_text_from_docx(file_bytes: bytes) -> str:
#     """Извлекает текст из Word документа."""
#     doc = docx.Document(io.BytesIO(file_bytes))
#     full_text = [para.text for para in doc.paragraphs]
#     return "\n".join(full_text)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Извлекает текст из Word документа, включая обычный текст и таблицы."""
    doc = docx.Document(io.BytesIO(file_bytes))
    full_text = []

    # 1. Сначала читаем обычные абзацы вне таблиц
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())

    # 2. Читаем текст из всех таблиц (Критически важно для вашего шаблона!)
    for table in doc.tables:
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                # Очищаем текст ячейки от лишних внутренних переносов строк
                clean_text = cell.text.strip().replace("\n", " ")
                if (
                    clean_text and clean_text not in row_data
                ):  # Избегаем дублей объединенных ячеек
                    row_data.append(clean_text)

            if row_data:
                # Склеиваем ячейки в одну строку через разделитель для ИИ
                full_text.append(" | ".join(row_data))

    return "\n".join(full_text)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Извлекает текст из PDF документа."""
    reader = PdfReader(io.BytesIO(file_bytes))
    text = ""
    for page in reader.pages:
        if page.extract_text():
            text += page.extract_text() + "\n"
    return text


def parse_document(file_bytes: bytes, filename: str) -> str:
    """Определяет формат и парсит документ."""
    if filename.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    elif filename.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif filename.endswith(".txt"):
        return file_bytes.decode("utf-8")
    else:
        raise ValueError("Формат не поддерживается. Загрузите PDF, Word или TXT.")
